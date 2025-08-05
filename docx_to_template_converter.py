import docx
import io
import re
import logging
import json
import os
from openai import OpenAI, RateLimitError, APIError
from spacy.tokens import Doc

# Configure logger
logger = logging.getLogger(__name__)

class QAValidationError(Exception):
    """Custom exception for non-fatal QA validation failures."""
    def __init__(self, message, issues):
        super().__init__(message)
        self.issues = issues

def _extract_text_from_docx(docx_stream: io.BytesIO) -> str:
    """Extracts all text from a .docx file stream."""
    try:
        doc = docx.Document(docx_stream)
    except KeyError as e:
        # This specific KeyError from python-docx indicates a malformed package.
        # It's often because the file is not a true .docx file (e.g., a .doc renamed).
        logger.error(f"Failed to parse DOCX file, it may be corrupted or not a valid .docx format. Error: {e}")
        raise ValueError(
            "The file is not a valid .docx file. It may be corrupted, an older format (.doc), "
            "or a different file type with a .docx extension."
        )

    full_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    docx_stream.seek(0)
    return full_text

def _get_semantic_map_from_llm(text: str, feedback_issues: list | None = None) -> dict:
    """
    Generates a structured semantic map by calling the OpenAI API.
    Optionally includes feedback from a previous failed attempt to guide the LLM.
    """
    logger.info("[Stage 1a] Generating semantic map via LLM.")
    if feedback_issues:
        logger.info(f"This is a retry. Incorporating {len(feedback_issues)} feedback points.")

    if not os.getenv("OPENAI_API_KEY"):
        raise ValueError("OPENAI_API_KEY environment variable is not set.")

    client = OpenAI()

    system_prompt = """
You are a world-class expert CV analyst and Jinja2 template engineer. Your work must be perfect and syntactically correct. Your task is to read the provided CV text and generate a structured JSON object to be used for converting the CV into a template.

**Your Goal:**
Create a JSON object with two main keys: `simple_replacements` and `block_replacements`.

1.  **`simple_replacements`**: A dictionary for simple text-for-placeholder swaps (e.g., `{"John Doe": "{{ name }}"}`).
2.  **`block_replacements`**: A list of objects for complex sections that need to be replaced with Jinja2 loops. Each object must have `original_block` and `new_block` keys.

**CRITICAL JINJA2 SYNTAX & LOGIC RULES:**

1.  **PERFECT LOOPS:** Every `{% for ... %}` loop MUST be closed with `{% endfor %}`. This is non-negotiable. Nested loops require nested `endfor` tags.
2.  **CORRECT LOOP VARIABLES:** Inside a loop, variables MUST be prefixed with the loop variable. For example, inside `{% for job in experiences %}`, you must use `{{ job.title }}`, NOT `{{ title }}`.
3.  **JOIN FILTERS:** For lists of skills, use the `| join(', ')` filter. Example: `{{ backend_technologies | join(', ') }}`.
4.  **NO INVENTED PLACEHOLDERS:** Only use the standard, logical placeholder names demonstrated in the example.

**Perfect Example of a `block_replacements` entry:**
-   **Input Text:** "Développeuse Web – fullstack\nCreative Web\nOctobre 2018\nOctobre 2025\nMISSIONS :\n- Task 1\n- Task 2"
-   **Required `new_block` code:**
    ```jinja2
{% for job in experiences %}
Développeuse Web – fullstack
{{ job.company }}
{{ job.start_date }} – {{ job.end_date }}

MISSIONS :
{% for task in job.tasks %}
- {{ task }}
{% endfor %}
{% endfor %}
    ```
    (Note the two `endfor` tags for the nested loops).

**Final Instruction:**
Before creating the final JSON, double-check all generated Jinja2 code in the `new_block` values for syntax errors. Your output must be flawless. Your output MUST be ONLY a valid JSON object.
"""
    user_prompt = f"Here is the CV text to be templated:\n\n```\n{text}\n```"

    if feedback_issues:
        feedback_prompt = (
            "IMPORTANT: This is a second attempt. Your previous attempt failed QA with the following errors. "
            "You MUST correct these specific errors in your new response. DO NOT repeat these mistakes.\n\n"
            "ERRORS TO FIX:\n"
        )
        for issue in feedback_issues:
            feedback_prompt += f"- {issue.get('error_type')}: {issue.get('description')}\n"

        user_prompt += f"\n\n{feedback_prompt}"

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.0,
            response_format={"type": "json_object"},
        )
        llm_output = response.choices[0].message.content
        semantic_map = json.loads(llm_output)
        logger.info("[Stage 1a] Successfully received structured semantic map from LLM.")
        return semantic_map
    except (RateLimitError, APIError) as e:
        logger.error(f"[Stage 1a] OpenAI API error: {e}")
        raise ValueError("The LLM service is currently unavailable or has hit a rate limit.")
    except Exception as e:
        logger.error(f"[Stage 1a] Failed to get semantic map from LLM: {e}", exc_info=True)
        raise

def stage1_get_semantic_map(text: str, use_llm: bool = True, feedback_issues: list | None = None) -> dict:
    """
    STAGE 1: Controller for generating the semantic map.
    """
    logger.info("[Stage 1] Starting semantic map generation.")
    if use_llm:
        return _get_semantic_map_from_llm(text, feedback_issues=feedback_issues)
    else:
        logger.warning("[Stage 1] Non-LLM method is not supported for this advanced templating. LLM is required.")
        return {"simple_replacements": {}, "block_replacements": []}


def _replace_text_in_paragraph(paragraph: 'docx.text.paragraph.Paragraph', replacements: dict[str, str]):
    """Helper function to perform simple search-and-replace on runs within a paragraph."""
    sorted_replacements = sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True)
    for old, new in sorted_replacements:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

def _delete_paragraph(paragraph):
    """Helper function to delete a paragraph from its parent."""
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    else:
        logger.warning("Attempted to delete a paragraph that has already been removed or orphaned. Skipping.")

def force_secure_loop_block(var_name: str, iterable: str, body: str, join: bool = False) -> str:
    """
    Deterministically builds a syntactically correct Jinja2 loop block.
    """
    if join:
        # Handles dictionary-style iteration for skills
        return f"""{{% for {var_name}, tools in {iterable}.items() %}}
{{{{ {var_name} }}}}: {{{{ tools | join(', ') }}}}
{{% endfor %}}"""
    # Handles standard list iteration
    return f"""{{% for {var_name} in {iterable} %}}
{body.strip()}
{{% endfor %}}"""

def inject_fallback_experience_block(doc):
    """
    Injects a fallback experience loop block at the end of the document
    if no valid experience block is found.
    """
    logger.warning("Injecting fallback experience block as it was missing.")
    doc.add_paragraph().add_run("\n--- Fallback Section ---").bold = True
    fallback_block = """\
{% for job in experiences %}
- {{ job.title }} at {{ job.company }} ({{ job.start_date }} - {{ job.end_date }})
  {{ job.description }}
{% endfor %}"""
    doc.add_paragraph(fallback_block)

def validate_and_finalize_template(doc, autofix: bool = True):
    """
    Validates the presence of required template structures, optionally
    injecting missing blocks if autofix is enabled.
    """
    logger.info("Performing final validation and finalization of template.")
    doc_text = "\n".join(p.text for p in doc.paragraphs)

    if "{{" not in doc_text and "{%" not in doc_text:
        raise ValueError("Validation failed: No Jinja2 syntax found in the document.")

    if doc_text.count("{% for") != doc_text.count("{% endfor %}"):
        raise SyntaxError("Validation failed: Unbalanced for/endfor blocks.")

    if "| join(" in doc_text and "{% for" not in doc_text:
        raise SyntaxError("Validation failed: `join` filter found outside of a loop context.")

    forbidden_edu_fields = ["education.year", "education.degree"]
    for field in forbidden_edu_fields:
        if f"{{{{ {field} }}}}" in doc_text:
            raise ValueError(f"Validation failed: Forbidden field used: {field}")

    # Check for experience section and autofix if enabled
    required_loop = "{% for job in experiences %}"
    if required_loop not in doc_text:
        if autofix:
            inject_fallback_experience_block(doc)
        else:
            raise ValueError(f"Validation failed: Missing required template logic: {required_loop}")

    logger.info("Final validation and finalization passed.")
    return True

def _normalize_text_for_match(text: str) -> str:
    """Aggressively normalizes text to ensure a match."""
    return re.sub(r'[\s\u00A0\t]+', ' ', text).strip()

def _replace_text_block(paragraphs: list, original_block: str, new_block: str):
    """
    Finds a multi-paragraph block of text within a list of paragraphs and replaces it.
    """
    logger.info(f"Attempting to replace block starting with: '{original_block.splitlines()[0] if original_block else ''}...'")
    original_lines = [line.strip() for line in original_block.splitlines() if line.strip()]
    normalized_original_block = " ".join(original_lines)
    normalized_original_block = _normalize_text_for_match(normalized_original_block)

    start_para_index = -1
    for i, p in enumerate(paragraphs):
        if _normalize_text_for_match(p.text) == _normalize_text_for_match(original_lines[0]):
            start_para_index = i
            break

    if start_para_index == -1:
        logger.warning("Could not find the starting paragraph of the block.")
        return False

    collected_paras, collected_text_normalized = [], []
    for i in range(start_para_index, len(paragraphs)):
        p = paragraphs[i]
        p_text_norm = _normalize_text_for_match(p.text)
        if not p_text_norm:
            continue
        collected_paras.append(p)
        collected_text_normalized.append(p_text_norm)

        current_block_normalized = " ".join(collected_text_normalized)
        if current_block_normalized == normalized_original_block:
            logger.info(f"Found matching block of {len(collected_paras)} paragraphs. Replacing.")
            first_para = collected_paras[0]
            first_para.clear()
            first_para.add_run(new_block)
            for para_to_delete in collected_paras[1:]:
                _delete_paragraph(para_to_delete)
            logger.info("Block replacement successful.")
            return True

    logger.warning("Could not find a matching block in the document to replace.")
    return False

def stage2_apply_annotations(docx_stream: io.BytesIO, semantic_map: dict) -> io.BytesIO:
    """
    STAGE 2: Applies annotations using a deterministic, rules-based approach.
    """
    logger.info("[Stage 2] Applying annotations deterministically.")
    try:
        doc = docx.Document(docx_stream)
    except Exception as e:
        logger.error(f"[Stage 2] Failed to load .docx file: {e}", exc_info=True)
        raise ValueError("Invalid or corrupted .docx file.")

    # Apply simple, direct replacements first
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    simple_replacements = semantic_map.get("simple_replacements", {})
    if simple_replacements:
        for paragraph in all_paragraphs:
            _replace_text_in_paragraph(paragraph, simple_replacements)

    # Deterministically build and replace complex blocks
    block_replacements = semantic_map.get("block_replacements", [])
    for block in block_replacements:
        original_text = block["original_block"]
        block_to_inject = ""

        # Rule-based identification of block type
        if "experience" in original_text.lower() or "mission" in original_text.lower():
            body = "{{ job.title }}\n{{ job.company }}\n{{ job.start_date }} - {{ job.end_date }}\n{{ job.description }}"
            block_to_inject = force_secure_loop_block("job", "experiences", body)
        elif "formation" in original_text.lower() or "education" in original_text.lower() or "diplôme" in original_text.lower():
            body = "{{ education.title }} - {{ education.institution }}"
            block_to_inject = force_secure_loop_block("education", "educations", body)
        elif "compétences" in original_text.lower() or "skills" in original_text.lower() or "outils" in original_text.lower():
             block_to_inject = force_secure_loop_block("category", "technical_skills", "", join=True)
        else:
            logger.warning(f"Could not determine block type for: '{original_text[:50]}...'. Skipping.")
            continue

        logger.debug(f"Injecting deterministic block for '{original_text[:20]}...':\n{block_to_inject}")

        current_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    current_paragraphs.extend(cell.paragraphs)
        _replace_text_block(current_paragraphs, original_text, block_to_inject)

    # Final validation and autofixing pass
    validate_and_finalize_template(doc, autofix=True)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    logger.info("[Stage 2] Deterministic annotation and validation successful.")
    return target_stream


def convert_docx_to_template(
    docx_stream: io.BytesIO,
    nlp_model: 'spacy.lang.en.English',
    use_llm: bool = True,
    feedback_issues: list | None = None
) -> io.BytesIO:
    """
    Orchestrates the multi-stage process of converting a DOCX to a Jinja2 template.
    Can accept feedback from a previous failed run to attempt self-correction.
    """
    logger.info("Starting multi-stage template conversion process.")
    full_text = _extract_text_from_docx(docx_stream)
    if not full_text.strip():
        logger.warning("No text found in the document. Returning original file.")
        return docx_stream

    # Stage 1: Get Semantic Map, potentially with feedback from a previous run
    semantic_map = stage1_get_semantic_map(full_text, use_llm=use_llm, feedback_issues=feedback_issues)
    if not semantic_map.get("simple_replacements") and not semantic_map.get("block_replacements"):
        logger.warning("No entities found to replace. Returning original document.")
        return docx_stream

    # Stage 2: Apply Annotations
    templated_docx_stream = stage2_apply_annotations(docx_stream, semantic_map)

    logger.info("Template conversion process completed successfully.")
    return templated_docx_stream
