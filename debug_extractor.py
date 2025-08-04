import sys
import docx

def extract_text_to_file(docx_path: str, output_path: str):
    """
    Extracts all paragraph text from a .docx file and saves it to a text file.
    """
    try:
        doc = docx.Document(docx_path)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("--- PARAGRAPHS FROM MAIN BODY ---\n")
            for i, p in enumerate(doc.paragraphs):
                f.write(f"Para {i}: {p.text}\n")

            f.write("\n\n--- PARAGRAPHS FROM TABLES ---\n")
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, p in enumerate(cell.paragraphs):
                            f.write(f"Table {t_idx}, Row {r_idx}, Cell {c_idx}, Para {p_idx}: {p.text}\n")

        print(f"Successfully extracted text to {output_path}")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python debug_extractor.py <path_to_docx_file>")
        sys.exit(1)

    docx_file = sys.argv[1]
    output_file = "extracted_text.txt"
    extract_text_to_file(docx_file, output_file)
