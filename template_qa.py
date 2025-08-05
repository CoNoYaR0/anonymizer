import io
import logging
import json
import docx
from openai import OpenAI, RateLimitError, APIError
from docx_to_template_converter import QAValidationError

# Configure logger
logger = logging.getLogger(__name__)

def _extract_text_from_docx(docx_stream: io.BytesIO) -> str:
    """Extracts all text from a .docx file stream."""
    doc = docx.Document(docx_stream)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    docx_stream.seek(0) # Reset stream for later use
    return full_text

def validate_template_with_llm(docx_stream: io.BytesIO) -> dict:
    """
    STAGE 3: Validates a generated .docx Jinja2 template using an LLM.
    """
    logger.info("[Stage 3] Starting LLM QA Review of the generated template.")

    template_text = _extract_text_from_docx(docx_stream)

    client = OpenAI()

    system_prompt = """
You are a meticulous QA engineer specializing in `docxtpl` Jinja2 templates. Your task is to validate the provided template text according to a strict set of rules.

**Input:** Text extracted from a .docx file that is supposed to be a Jinja2 template.

**Output:** A JSON object with two keys:
1.  `is_valid` (boolean): `true` if the template is perfect, `false` otherwise.
2.  `issues` (array of objects): A list of any issues found. If `is_valid` is `true`, this should be an empty array.

**Validation Rules:**

1.  **Jinja2 Syntax:** Check for any syntax errors.
    - All tags must be correctly opened and closed (e.g., `{{ variable }}`, `{% for item in items %}`).
    - `for` loops must have a corresponding `endfor`.
    - No mismatched or unclosed brackets.

2.  **Placeholder Naming Conventions:** All placeholders must follow these specific naming conventions:
    - **Header:** `{{ user_initials }}`, `{{ job_title }}`, `{{ years_of_experience }}`, `{{ current_client }}`
    - **Education:** Must be in a loop `{% for education in educations %}` with `{{ education.title }}`, `{{ education.institution }}`, etc.
    - **Skills:** Should use `| join(', ')` for lists, e.g., `{{ backend_technologies | join(', ') }}`.
    - **Experience:** Must be in a loop `{% for job in experiences %}` with `{{ job.title }}`, `{{ job.company }}`, etc.

3.  **No Unintended Replacements:**
    - Ensure that common words or technologies have not been accidentally converted into placeholders (e.g., "Redis" should not be `{{ name }}`).
    - The template should look clean and professional.

**Example of a valid response for a faulty template:**
```json
{
  "is_valid": false,
  "issues": [
    {
      "error_type": "Syntax Error",
      "description": "Found an unclosed Jinja2 tag: '{{ user_initials'."
    },
    {
      "error_type": "Naming Convention Violation",
      "description": "Found '{{ company }}' which should be '{{ job.company }}' inside the experience loop."
    }
  ]
}
```
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Please validate the following template text:\n\n```\n{template_text}\n```"},
            ],
            temperature=0.0, # Be very strict and deterministic
            response_format={"type": "json_object"},
        )
        llm_output = response.choices[0].message.content
        validation_result = json.loads(llm_output)

        if validation_result.get("is_valid"):
            logger.info("[Stage 3] LLM QA Review PASSED.")
            return {"is_valid": True, "issues": []}
        else:
            issues = validation_result.get('issues', [])
            logger.warning(f"[Stage 3] LLM QA Review FAILED. Raising QAValidationError with {len(issues)} issues.")
            raise QAValidationError("LLM QA validation failed.", issues=issues)

    except (RateLimitError, APIError) as e:
        logger.error(f"[Stage 3] OpenAI API error during validation: {e}")
        # In case of validation failure, it's safer to assume the template is invalid.
        return {"is_valid": False, "issues": [{"error_type": "API Error", "description": "The validation service is currently unavailable."}]}
    except Exception as e:
        logger.error(f"[Stage 3] Failed to validate template with LLM: {e}", exc_info=True)
        return {"is_valid": False, "issues": [{"error_type": "Exception", "description": "An unexpected error occurred during the validation process."}]}
