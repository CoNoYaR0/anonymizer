import pytest
from fastapi.testclient import TestClient
import docx
import io

# Import the FastAPI app instance from your main application file
from main import app

# Create a client for testing
client = TestClient(app)

@pytest.fixture
def sample_docx_stream() -> io.BytesIO:
    """
    Creates a sample .docx file in memory with text that matches
    the mock_structured_data.
    """
    document = docx.Document()
    document.add_paragraph("Jean Dupont") # "name"
    document.add_paragraph("Développeur") # "job_title"
    document.add_paragraph("Contact: jean.dupont@email.com") # "contact.email"

    # Add the block that will be replaced
    document.add_paragraph("Développeuse Web – fullstack")
    document.add_paragraph("Creative Web")
    document.add_paragraph("Octobre 2018 - Octobre 2025")
    document.add_paragraph("MISSIONS :")
    document.add_paragraph("Refonte complète et maintenance d’un CMS")
    document.add_paragraph("Développement backend d’API REST")

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream

def test_convert_to_template_endpoint_success(sample_docx_stream: io.BytesIO, monkeypatch):
    """
    Tests the /convert-to-template endpoint with a valid .docx file,
    mocking the LLM call to ensure deterministic results.
    """
    # 1. Define the mock semantic map that the LLM would return
    mock_semantic_map = {
        "simple_replacements": {
            "Jean Dupont": "{{ name }}",
            "Développeur": "{{ title }}",
            "jean.dupont@email.com": "{{ email }}"
        },
        "block_replacements": [
            {
                "original_block": "Développeuse Web – fullstack\nCreative Web\nOctobre 2018 - Octobre 2025\nMISSIONS :\nRefonte complète et maintenance d’un CMS\nDéveloppement backend d’API REST",
                "new_block": "{% for job in experiences %}{{ job.title }}{% endfor %}" # This will be deterministically replaced
            }
        ]
    }

    def mock_get_semantic_map(text: str, feedback_issues: list | None = None):
        # The mock doesn't need to use the feedback, just accept it
        return mock_semantic_map

    # 2. Apply the mocks
    monkeypatch.setattr(
        "docx_to_template_converter._get_semantic_map_from_llm",
        mock_get_semantic_map
    )
    # We also mock the final QA stage to isolate Stage 1 and 2
    monkeypatch.setattr(
        "template_qa.validate_template_with_llm",
        lambda docx_stream: {"is_valid": True, "issues": []}
    )

    # 3. Run the test with the mock in place
    files = {'file': ('sample_cv.docx', sample_docx_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    response = client.post("/convert-to-template", files=files)

    # --- Assertions ---
    assert response.status_code == 200
    assert 'attachment; filename="sample_cv_template.docx"' in response.headers['content-disposition']

    # Read the content of the response
    response_stream = io.BytesIO(response.content)
    returned_doc = docx.Document(response_stream)
    returned_text = "\n".join([p.text for p in returned_doc.paragraphs])
    print(f"Returned text:\n---\n{returned_text}\n---")

    # Check for simple replacements that were in the mock
    assert "{{ name }}" in returned_text
    assert "{{ title }}" in returned_text
    assert "Contact: {{ email }}" in returned_text

    # Check for the deterministically generated block replacement
    assert "{% for job in experiences %}" in returned_text
    assert "{{ job.title }}" in returned_text
    assert "{{ job.company }}" in returned_text
    assert "{{ job.description }}" in returned_text
    assert "{% endfor %}" in returned_text

    # Check that original data is gone
    assert "Jean Dupont" not in returned_text
    assert "Développeur" not in returned_text
    assert "Creative Web" not in returned_text

def test_convert_to_template_endpoint_invalid_file_type():
    """
    Tests the /convert-to-template endpoint with an invalid file type.
    """
    # Create a dummy text file
    txt_stream = io.BytesIO(b"this is not a docx file")
    files = {'file': ('test.txt', txt_stream, 'text/plain')}

    response = client.post("/convert-to-template", files=files)

    assert response.status_code == 400
    assert response.json() == {"error": "Invalid file type. Please upload a valid .docx file."}
