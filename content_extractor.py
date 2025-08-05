import io
import logging
import json
from typing import IO
from openai import OpenAI
import pytesseract
from pdf2image import convert_from_bytes
import docx

# Configure logger
logger = logging.getLogger(__name__)

def _get_text_from_pdf(file_stream: IO[bytes]) -> str:
    """Extracts text from a PDF file stream using OCR."""
    logger.info("Extracting text from PDF using OCR...")
    try:
        pdf_bytes = file_stream.read()
        images = convert_from_bytes(pdf_bytes)
        text = ""
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang='fra')
            text += page_text + "\n"
            logger.debug(f"Extracted text from page {i+1}.")
        logger.info("PDF text extraction successful.")
        return text
    except Exception as e:
        logger.error(f"OCR processing failed: {e}", exc_info=True)
        raise ValueError("Failed to extract text from PDF. Ensure Tesseract and Poppler are installed.")

def _get_text_from_docx(file_stream: IO[bytes]) -> str:
    """Extracts text from a DOCX file stream."""
    logger.info("Extracting text from DOCX...")
    try:
        doc = docx.Document(file_stream)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        logger.info("DOCX text extraction successful.")
        return full_text
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}", exc_info=True)
        raise ValueError("Failed to process .docx file. It may be corrupted.")

def extract_json_from_cv(file_stream: IO[bytes], content_type: str) -> dict:
    """
    Orchestrates the extraction of a structured JSON object from a CV file.

    Args:
        file_stream: The byte stream of the uploaded CV.
        content_type: The MIME type of the file.

    Returns:
        A dictionary containing the structured CV data.
    """
    raw_text = ""
    if content_type == "application/pdf":
        raw_text = _get_text_from_pdf(file_stream)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raw_text = _get_text_from_docx(file_stream)
    elif content_type.startswith("image/"):
        logger.info("Extracting text from image using OCR...")
        # For images, pytesseract can read the bytes directly
        try:
            image_bytes = file_stream.read()
            raw_text = pytesseract.image_to_string(image_bytes, lang='fra')
            logger.info("Image text extraction successful.")
        except Exception as e:
            logger.error(f"Image OCR processing failed: {e}", exc_info=True)
            raise ValueError("Failed to extract text from image.")
    else:
        raise ValueError(f"Unsupported file type: {content_type}")

    if not raw_text.strip():
        raise ValueError("No text could be extracted from the document.")

    # --- LLM Data Structuring ---
    logger.info("Sending extracted text to LLM for structuring into JSON.")
    client = OpenAI()
    system_prompt = """
You are a world-class HR data analyst. Your task is to read the raw text from a CV and convert it into a structured, clean JSON object. The JSON keys must be in English.

**Your Goal:**
Create a JSON object with the following schema. All fields are optional, but you must extract as much information as you can accurately.
- `name` (string)
- `title` (string)
- `email` (string)
- `phone` (string)
- `location` (string)
- `summary` (string)
- `experiences` (array of objects):
  - `title` (string)
  - `company` (string)
  - `start_date` (string)
  - `end_date` (string)
  - `description` (string)
- `educations` (array of objects):
  - `title` (string)
  - `institution` (string)
  - `date` (string)
- `skills` (object of arrays): where keys are categories (e.g., "Languages", "Backend", "Databases") and values are lists of skills.

Your output MUST be ONLY a valid JSON object. Do not include any commentary.
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Here is the raw CV text:\n\n```\n{raw_text}\n```"},
            ],
            temperature=0.0,
            response_format={"type": "json_object"},
        )
        structured_data = json.loads(response.choices[0].message.content)
        logger.info("Successfully received structured JSON from LLM.")
        return structured_data
    except Exception as e:
        logger.error(f"LLM data structuring failed: {e}", exc_info=True)
        raise ValueError("Failed to structure data using the LLM.")
