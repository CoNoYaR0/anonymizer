import logger_config # Import to configure logging
import logging
from dotenv import load_dotenv, find_dotenv

# Load environment variables from .env file BEFORE any other imports that might
# depend on them (like llm_refiner). find_dotenv() will locate the .env file
# in the project root.
load_dotenv(find_dotenv())

import os
import re
import uuid
import hashlib
from fastapi import FastAPI, File, UploadFile, HTTPException, Body
from pydantic import BaseModel, Field
import pytesseract
from pdf2image import convert_from_bytes
import spacy
from supabase import create_client, Client
import io
from PIL import Image
import docx
from docx.shared import Inches
import psutil
from llm_refiner import refine_extraction_with_llm

# Get a logger for the current module
logger = logging.getLogger(__name__)

# --- Environment-based Configuration ---
DEBUG_MODE = os.getenv("DEBUG", "False").lower() in ("true", "1", "t")

class ExtractedEntities(BaseModel):
    persons: list[str] = Field(..., example=["Jean Dupont"])
    locations: list[str] = Field(..., example=["Paris"])
    emails: list[str] = Field(..., example=["jean.dupont@example.com"])
    phones: list[str] = Field(..., example=["01 23 45 67 89"])
    skills: list[str] = Field(..., example=["Python", "FastAPI"])
    experience: list[str] = Field(..., example=["Développeur chez Example Corp"])

class AnonymizeRequest(BaseModel):
    filename: str
    entities: ExtractedEntities
    raw_text: str

# Load spaCy model
try:
    nlp = spacy.load("fr_core_news_lg")
except OSError:
    print("Downloading spaCy model 'fr_core_news_lg'...")
    from spacy.cli import download
    download("fr_core_news_lg")
    nlp = spacy.load("fr_core_news_lg")


app = FastAPI(title="CV Anonymizer API")

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")

if not SUPABASE_URL or not SUPABASE_ANON_KEY:
    print("Warning: Supabase credentials not found. Supabase integration will be disabled.")
    supabase: Client | None = None
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)

@app.get("/")
def read_root():
    """Root endpoint to check if the API is running."""
    return {"message": "Welcome to the CV Anonymizer API!"}

@app.get("/status")
def get_status():
    """Returns the current status of the server, focusing on the app's resource usage."""
    process = psutil.Process(os.getpid())

    # Get process-specific memory usage
    memory_info = process.memory_info()
    memory_used_mb = memory_info.rss / (1024 ** 2)  # rss is typically the most relevant metric

    # Get overall system disk usage (disk usage is not process-specific)
    disk_info = psutil.disk_usage('/')

    return {
        "app_cpu_usage_percent": process.cpu_percent(interval=0.1),
        "app_memory_usage": {
            "used_mb": f"{memory_used_mb:.2f} MB"
        },
        "system_disk_usage": {
            "total": f"{disk_info.total / (1024**3):.2f} GB",
            "used": f"{disk_info.used / (1024**3):.2f} GB",
            "free": f"{disk_info.free / (1024**3):.2f} GB",
            "percent": disk_info.percent
        }
    }

@app.post("/upload")
async def upload_cv(file: UploadFile = File(...)):
    """
    Uploads a CV, extracts text, refines it with an LLM, and saves the result.
    """
    logger.info(f"Received file '{file.filename}' for upload.")
    if file.content_type != "application/pdf":
        logger.warning(f"Invalid file type '{file.content_type}' received.")
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF.")

    try:
        pdf_bytes = await file.read()

        # --- Deduplication Check ---
        file_hash = hashlib.sha256(pdf_bytes).hexdigest()
        logger.info(f"Calculated SHA256 hash for '{file.filename}': {file_hash}")

        if supabase:
            try:
                logger.debug(f"Checking for existing hash in the database...")
                existing_extraction = supabase.table("extractions").select("id").eq("file_hash", file_hash).single().execute()
                if existing_extraction.data:
                    extraction_id = existing_extraction.data['id']
                    logger.info(f"Duplicate file detected. Found existing extraction_id: {extraction_id} for hash: {file_hash}")
                    return {"extraction_id": extraction_id, "detail": "File already processed."}
            except Exception as e:
                # If the query fails for some reason (e.g., PostgREST error), log it but continue processing.
                # It's safer to re-process than to fail the upload.
                logger.error(f"Could not check for existing file hash. Proceeding with processing. Error: {e}")

        logger.info("Starting OCR processing...")
        try:
            images = convert_from_bytes(pdf_bytes)
            text = ""
            for i, img in enumerate(images):
                page_text = pytesseract.image_to_string(img, lang='fra')
                text += page_text + "\n"
                logger.debug(f"Extracted text from page {i+1}.")
            logger.info("OCR processing completed.")
        except Exception as ocr_error:
            logger.error(f"OCR processing failed: {ocr_error}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"OCR processing failed. Ensure Tesseract and Poppler are installed.")

        if not text.strip():
            logger.warning("No text could be extracted from the PDF.")
            raise HTTPException(status_code=400, detail="Could not extract any text from the PDF.")

        logger.info("Starting initial data extraction with spaCy...")
        doc = nlp(text)
        persons = list(set([ent.text for ent in doc.ents if ent.label_ == "PER"]))
        locations = list(set([ent.text for ent in doc.ents if ent.label_ == "LOC"]))
        emails = list(set(re.findall(r'[\w\.-]+@[\w\.-]+', text)))
        phones = list(set(re.findall(r'(\d{2}[-\.\s]?){4}\d{2}', text)))
        initial_extraction = {"persons": persons, "locations": locations, "emails": emails, "phones": phones, "skills": [], "experience": []}
        logger.info("Initial data extraction completed.")

        # --- LLM Refinement ---
        llm_success, llm_result = refine_extraction_with_llm(text, initial_extraction)
        if not llm_success:
            logger.error("LLM refinement failed.")
            error_detail = "LLM refinement failed. Upstream service may be unavailable."
            if DEBUG_MODE:
                error_detail = llm_result
            raise HTTPException(status_code=502, detail=error_detail)

        refined_entities = llm_result

        final_data_to_save = {"filename": file.filename, "entities": refined_entities, "raw_text": text}
        logger.info("Successfully prepared final data object.")

        # --- Supabase Integration ---
        if supabase:
            safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
            file_path = f"uploads/{uuid.uuid4()}_{safe_filename}"
            logger.info(f"Uploading file to Supabase at path: {file_path}")
            try:
                supabase.storage.from_("cvs").upload(file_path, pdf_bytes, {"content-type": "application/pdf"})
                logger.info("File uploaded to Supabase Storage.")

                logger.info("Inserting new extraction data into Supabase DB...")
                db_response = supabase.table("extractions").insert({
                    "filename": file.filename,
                    "storage_path": file_path,
                    "data": final_data_to_save,
                    "file_hash": file_hash  # Add the hash to the new record
                }).execute()
                new_extraction_id = db_response.data[0]['id']
                logger.info(f"New extraction data saved with ID: {new_extraction_id}")

                return {"extraction_id": new_extraction_id}
            except Exception as e:
                logger.error(f"Supabase operation failed: {e}", exc_info=True)
                raise HTTPException(status_code=500, detail="Failed to save data to backend storage.")
        else:
            logger.warning("Supabase is not configured. Skipping database operations.")
            # In a real scenario, you might want to block this or handle it differently.
            # For now, we'll return the data without saving, which is not ideal.
            return {"extraction_id": None, "detail": "Supabase not configured", "data": final_data_to_save}

    except HTTPException:
        # Re-raise HTTPException to let FastAPI handle it
        raise
    except Exception as e:
        logger.critical(f"An unexpected error occurred during file processing for '{file.filename}': {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="An unexpected error occurred during file processing.")


def anonymize_text(text: str, entities: ExtractedEntities) -> str:
    """Replaces personal information in the text with anonymized placeholders."""

    # Anonymize persons to initials
    for person in entities.persons:
        initials = "".join([name[0].upper() for name in person.split()])
        text = text.replace(person, f"Person ({initials})")

    # Anonymize emails
    for email in entities.emails:
        text = text.replace(email, "[EMAIL REDACTED]")

    # Anonymize phones
    for phone in entities.phones:
        text = text.replace(phone, "[PHONE REDACTED]")

    # Anonymize locations
    for location in entities.locations:
        text = text.replace(location, "[LOCATION REDACTED]")

    return text


@app.get("/anonymize/{extraction_id}")
async def anonymize_cv_by_id(extraction_id: int):
    """
    Fetches an extraction by ID, anonymizes it, generates a DOCX file,
    and returns a secure download link.
    """
    logger.info(f"Received request to anonymize extraction ID: {extraction_id}")
    if not supabase:
        logger.error("Anonymization request failed: Supabase is not configured.")
        raise HTTPException(status_code=503, detail="Supabase is not configured. Cannot fetch extraction data.")

    try:
        logger.info(f"Fetching extraction data for ID: {extraction_id} from Supabase.")
        response = supabase.table("extractions").select("*").eq("id", extraction_id).single().execute()
        extraction_data = response.data
        logger.info(f"Successfully fetched data for extraction ID: {extraction_id}.")

        if not extraction_data:
            logger.warning(f"Extraction ID: {extraction_id} not found in the database.")
            raise HTTPException(status_code=404, detail="Extraction not found.")

        request_entities = ExtractedEntities(**extraction_data['data']['entities'])
        request = AnonymizeRequest(
            filename=extraction_data['data']['filename'],
            entities=request_entities,
            raw_text=extraction_data['data']['raw_text']
        )
        logger.debug("Successfully reconstructed request data from fetched record.")

        logger.info("Anonymizing text...")
        anonymized_text = anonymize_text(request.raw_text, request.entities)
        logger.info("Text anonymization complete.")

        logger.info("Generating DOCX document...")
        doc = docx.Document()
        doc.add_heading('CV Anonymisé', 0)
        doc.add_paragraph(anonymized_text)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        logger.info("DOCX document generated successfully.")

        anonymized_filename = f"anonymized_{request.filename.replace('.pdf', '.docx')}"
        file_path = f"anonymized_cvs/{uuid.uuid4()}_{anonymized_filename}"

        logger.info(f"Uploading anonymized DOCX to Supabase at path: {file_path}")
        supabase.storage.from_("cvs").upload(
            file_path,
            doc_io.getvalue(),
            {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )
        logger.info("Anonymized file uploaded to Supabase storage.")

        logger.info("Creating signed download URL...")
        # 24-hour expiry
        download_url_response = supabase.storage.from_("cvs").create_signed_url(file_path, 60 * 60 * 24)
        logger.info("Signed URL created successfully.")

        return {"download_url": download_url_response['signedURL']}

    except HTTPException:
        # Re-raise HTTPException to let FastAPI handle it
        raise
    except Exception as e:
        logger.critical(f"An unexpected error occurred during anonymization for ID {extraction_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during anonymization.")
