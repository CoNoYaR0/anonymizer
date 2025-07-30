import os
import json
import re
import spacy
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Load the French spaCy model
nlp = spacy.load("fr_core_news_md")

def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file using python-docx.
    """
    logging.info(f"Attempting to extract text from: {docx_path}")
    try:
        doc = Document(docx_path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        logging.info("Successfully extracted text from DOCX file.")
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return f"Error extracting text from DOCX: {e}"

def extract_contact_info(text):
    """
    Extracts contact information (email and phone number) from the text.
    """
    email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    phone = re.search(r'(\d{2}[-\.\s]?){4}\d{2}', text)
    return email.group(0) if email else None, phone.group(0) if phone else None

def extract_name(text):
    """
    Extracts the name from the text using spaCy's NER.
    """
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PER":
            return ent.text
    return None

def pseudonymize_name(name):
    """
    Pseudonymizes the name by taking the initials.
    """
    if not name:
        return None
    return "".join([n[0] for n in name.split()]).upper()

def extract_sections(text):
    """
    Extracts sections (experiences, competencies, education, technologies) from the text.
    This is a simplified implementation and may need to be improved.
    """
    sections = {
        "experiences": [],
        "competencies": [],
        "education": [],
        "technologies": []
    }

    # This is a placeholder for the actual implementation
    # A more robust solution would use regex or NLP to identify sections and their content

    lines = text.split('\n')
    current_section = None

    for line in lines:
        if "expérience" in line.lower():
            current_section = "experiences"
        elif "compétence" in line.lower():
            current_section = "competencies"
        elif "formation" in line.lower() or "cursus" in line.lower():
            current_section = "education"
        elif "technologies" in line.lower():
            current_section = "technologies"
        elif current_section and line.strip():
            sections[current_section].append(line.strip())

    return sections

def create_docx(output_path, data):
    """
    Creates a DOCX file with the provided data.
    """
    try:
        doc = Document()
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return False

def main():
    """
    Main function to run the anonymization process.
    """
    # Create a new, empty docx file to use as a template
    doc = Document()
    doc.save("templates/template_cv.docx")

    # File paths
    docx_path = "templates/template_cv.docx"
    output_dir = "outputs"

    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. Extract text from DOCX
    text = extract_text_from_docx(docx_path)
    if text.startswith("Error"):
        print(text)
        return

    # 2. Extract information
    name = extract_name(text)
    email, phone = extract_contact_info(text)
    sections = extract_sections(text)

    # 3. Pseudonymize
    candidate_id = pseudonymize_name(name)

    # 4. Build JSON
    data = {
        "candidate_id": candidate_id,
        "contact": {
            "email": "email_001@example.com",
            "phone": "phone_001"
        },
        "experiences": sections["experiences"],
        "competencies": sections["competencies"],
        "education": sections["education"],
        "technologies": sections["technologies"]
    }

    # Save JSON
    json_output_path = os.path.join(output_dir, f"anonymized_cv_{candidate_id}.json")
    with open(json_output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

    # 5. Create DOCX
    docx_output_path = os.path.join(output_dir, f"anonymized_cv_{candidate_id}.docx")

    # Data for the docx
    docx_data = {
        "candidate_id": candidate_id,
        "email": "email_001@example.com",
        "phone": "phone_001",
        "experiences": "\n".join(sections["experiences"]),
        "competencies": "\n".join(sections["competencies"]),
        "education": "\n".join(sections["education"]),
        "technologies": "\n".join(sections["technologies"]),
    }

    if not create_docx(docx_output_path, docx_data):
        print("Failed to create DOCX file.")

    # Report missing fields
    errors = []
    if not name:
        errors.append("Name not found.")
    if not email:
        errors.append("Email not found.")
    if not phone:
        errors.append("Phone number not found.")

    if errors:
        with open(os.path.join(output_dir, "errors.txt"), "w") as f:
            for error in errors:
                f.write(f"{error}\n")

if __name__ == "__main__":
    main()
