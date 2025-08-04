from docx import Document
import re

# --- Configuration ---
SOURCE_DOC_PATH = "templates/Dossier_competences_OHA.docx"
TEMPLATE_SAVE_PATH = "templates/cv_template.docx"
COMMERCIAL_SECTION_HEADING = "CONTACT COMMERCIAL" # Assumption: This text marks the start of the static section.

# This dictionary maps the exact text from the source document to the Jinja2 placeholder.
# This requires careful manual setup based on the content of `Dossier_competences_OHA.docx`.
# NOTE: This is a sample mapping. The actual text from the document must be used.
REPLACEMENT_MAP = {
    "Khalil Beji": "{{ initiales }}",
    "Consultant BI": "{{ experiences[0].job_title }}",
    "Mission au sein de La Banque Postale": "Mission au sein de {{ experiences[0].company_name }}",
    # ... other personal data fields

    # Experience 1
    "La Banque Postale": "{{ experiences[0].company_name }}",
    "Février 2022 - en cours": "{{ experiences[0].start_date }} - {{ experiences[0].end_date }}",
    "Contexte : Dans le cadre...": "{{ experiences[0].job_context }}",
    "Rédaction des spécifications fonctionnelles": "{{ experiences[0].missions[0] }}",
    # ... other missions for experience 1
    "BO, DWH, SQL": "{{ experiences[0].technologies | join(', ') }}",

    # Experience 2
    "Crédit Agricole": "{{ experiences[1].company_name }}",
    "Janvier 2021 - Février 2022": "{{ experiences[1].start_date }} - {{ experiences[1].end_date }}",
    # ... and so on for all experiences

    # Skills - This part is more complex and might require looping in the template
    # For the one-time creation, we can replace a whole block.
    "Langages : SQL, PL/SQL, Python": "{% for skill in skills %}{% if skill.category == 'Langages' %}{{ skill.skills_list | join(', ') }}{% endif %}{% endfor %}",
    "Bases de données : Teradata, Oracle, PostgreSQL": "{% for skill in skills %}{% if skill.category == 'Bases de données' %}{{ skill.skills_list | join(', ') }}{% endif %}{% endfor %}",
    # ... and so on for all skill categories
}

def is_in_commercial_section(paragraph: 'Paragraph') -> bool:
    """
    Checks if a paragraph is under the commercial section heading.
    This is a simple implementation. A more robust one might use flags.
    """
    # This assumes the commercial section is at the end. We check all previous paragraphs.
    # This is computationally inefficient but simple and fine for a one-time script.
    # A better way is to set a flag when the heading is found.
    current_element = paragraph.element
    while current_element is not None:
        if current_element.tag.endswith('p'):
            p = current_element
            text = "".join(run.text for run in p.iter_runs())
            if COMMERCIAL_SECTION_HEADING.lower() in text.lower():
                return True
        current_element = current_element.getprevious()
    return False


def replace_text_in_paragraph(paragraph):
    """Replaces text in a paragraph, preserving formatting."""
    for key, value in REPLACEMENT_MAP.items():
        if key in paragraph.text:
            # This is a simple replacement and may not preserve complex formatting perfectly.
            # A more advanced version would iterate through runs.
            inline = paragraph.runs
            # Replace strings and retain formatting
            for i in range(len(inline)):
                if key in inline[i].text:
                    text = inline[i].text.replace(key, str(value))
                    inline[i].text = text


def create_template():
    """
    Loads the source DOCX, replaces content with Jinja2 placeholders,
    and saves it as a new template file.
    """
    print(f"Loading source document: {SOURCE_DOC_PATH}")
    try:
        document = Document(SOURCE_DOC_PATH)
    except Exception as e:
        print(f"\n[ERROR] Could not open source document at '{SOURCE_DOC_PATH}'.")
        print("Please make sure you have placed the 'Dossier_competences_OHA.docx' file in the 'templates' directory.")
        return

    print("Processing document to create template...")

    in_commercial_section_flag = False

    # Process paragraphs
    for para in document.paragraphs:
        # Check for the commercial section heading
        if COMMERCIAL_SECTION_HEADING.lower() in para.text.lower():
            in_commercial_section_flag = True
            print("Found 'CONTACT COMMERCIAL' section. Skipping further replacements.")

        if not in_commercial_section_flag:
            replace_text_in_paragraph(para)

    # Process tables (many CVs use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if COMMERCIAL_SECTION_HEADING.lower() in para.text.lower():
                        in_commercial_section_flag = True
                        print("Found 'CONTACT COMMERCIAL' section in a table. Skipping further replacements.")

                    if not in_commercial_section_flag:
                        replace_text_in_paragraph(para)

    print(f"Saving template to: {TEMPLATE_SAVE_PATH}")
    document.save(TEMPLATE_SAVE_PATH)
    print("Template creation complete.")


if __name__ == "__main__":
    create_template()
